from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import hashlib
import os
import re
from pathlib import Path
from typing import List

from docx import Document
from pypdf import PdfReader


@dataclass(frozen=True)
class Chunk:
    text: str
    source: str
    chunk_id: int


def file_sha256(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    return hashlib.sha256(data).hexdigest()


def save_uploaded_file(uploaded_file, folder: str = "uploads") -> Path:
    """Persist an uploaded file to disk for traceability/debugging."""
    out_dir = Path(folder)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", uploaded_file.name)
    path = out_dir / safe_name
    path.write_bytes(uploaded_file.getvalue())
    return path


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(uploaded_file) -> str:
    source = BytesIO(uploaded_file.getvalue()) if hasattr(uploaded_file, "getvalue") else uploaded_file
    reader = PdfReader(source)
    parts: List[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return _clean_text("\n".join(parts))


def extract_text_from_docx(uploaded_file) -> str:
    doc = Document(BytesIO(uploaded_file.getvalue()))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return _clean_text("\n".join(parts))


def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    if name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 180) -> List[str]:
    text = _clean_text(text)
    if not text:
        return []

    chunks: List[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == n:
            break
        start = max(0, end - overlap)
    return chunks


def split_with_metadata(text: str, source: str) -> List[Chunk]:
    chunks = chunk_text(text)
    return [Chunk(text=c, source=source, chunk_id=i + 1) for i, c in enumerate(chunks)]
